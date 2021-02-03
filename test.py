import pdfplumber
import pandas as pd
pdffile = './test.pdf'
pdf = pdfplumber.open(pdffile)
p0 = pdf.pages[0]
text = p0.extract_text()
target='.  '
num_lis=[0]
prob_num_lis=[]
prob_lis=[]
#exp_num = input("Expected num of prob: ")
exp_num=70

for i in range(0,int(exp_num)):
    prob_num_lis.append(0)
    prob_lis.append(0)


prob=[]
page_lim_lis=[]
num = 0
prob_num=0
#print(len(pdf.pages))
for pages_num in range(len(pdf.pages)):
    text = pdf.pages[pages_num].extract_text()
    #print(f"Dealing with page{pages_num+1}...")
    page_lim_lis.append(len(text))
    page_lim = len(text)
    pos = 0
    pos_lis=[]
    subtext = text
    while(pos<len(text)):
        if target in subtext:
            pos = text.index(target, pos+1)
            #print(pos)
            
            if num>=9:
                key = text[pos-2:pos]
                if key.isdigit():
                    #prob_num_lis.append(key)
                    key = int(key)
                    prob_num_lis[key-1]=key
                    num+=1
                    pos_lis.append(pos)

            else:
                key = text[pos-1:pos]
                if key.isdigit():
                    key = int(key)
                    #prob_num_lis.append(key)
                    prob_num_lis[key-1]=key
                    num+=1
                    pos_lis.append(pos)
            subtext = text[pos+2:]
        else:
            pos = len(text)
    #print(f"For page{pages_num+1}, pos={pos_lis}")
    for i in range(len(pos_lis)):
        if i != (len(pos_lis)-1):
            problem = text[pos_lis[i]:pos_lis[i+1]]
        else:
            problem = text[pos_lis[i]:page_lim]
        prob_lis[prob_num]=problem
        prob_num+=1



missing = [i for i,x in enumerate(prob_num_lis) if x==0]

for i in missing:
    print(f"Missing Prob num. {i+1}, please check it manually!!!")
    prob_lis = prob_lis[:i]+['0']+prob_lis[i:]
    prob_lis.pop()
missing_prob = [i for i,x in enumerate(prob_lis) if x==0]
for i in missing_prob:
    print(f"Missing Prob. {i+1}, please check it manually!!!")




chapter = ['酸鹼','氧化還原','沉澱','氣體','平衡','反應速率',
'反應熱','電化學','無機','有機','週期表','電子軌域','原子構造','溶液','基本化學','濃度','未分類']


acid_keyword = ['酸鹼','H+','OH-','pH','pOH','Ka','Kw','酸性','中性','鹼性','離子積','滴定','共軛酸鹼','弱鹼','弱酸','強酸','強鹼','氫離子','氫氧根離子']
oxred_keyword = ['氧化','還原']
precipitaiton_keyword = ['沉澱','沈澱','ksp']
gas_keyword =['氣體','理想氣體','波以耳','查理','給呂薩克','nRT','氣壓計','大氣壓','分壓']
balance_keyword =['平衡','平衡常數','K','勒沙特列']
rate_keyword=['反應速率','速率定律式']
heat_keyword=['反應熱','生成熱','分解熱','燃燒熱']
electric_keyword=['電化學','電位','電壓','電流','庫侖']
inorganic_keyword=['無機','錯合物','錯離子','配位']
organic_keyword=['有機','烷','烯','炔','醯胺','胺','異構物','醇','醛','酮','酯','醚']
periodic_keyword=['週期表','電負度','電子親和力','游離能','原子半徑','金屬性']
orbital_keyword=['軌域','混成','共價鍵','價鍵理論','金屬鍵','離子鍵','氫鍵','凡得瓦力','鍵角','鍵長','鍵能','極性','分子形狀','共振','單鍵','雙鍵','參鍵']
atom_keyword=['原子','質子','電子','中子','原子序','質量數']
solution_keyword=['溶液','拉午耳','蒸氣壓','理想溶液']
basic_keyword=['係數平衡']
concentration_keyword=['溶解度','濃度','飽和','不飽和','過飽和']

keyword_prob = []
keyword_prob.append(acid_keyword)
keyword_prob.append(oxred_keyword)
keyword_prob.append(precipitaiton_keyword)
keyword_prob.append(gas_keyword)
keyword_prob.append(balance_keyword)
keyword_prob.append(rate_keyword)
keyword_prob.append(heat_keyword)
keyword_prob.append(electric_keyword)
keyword_prob.append(inorganic_keyword)
keyword_prob.append(organic_keyword)
keyword_prob.append(periodic_keyword)
keyword_prob.append(orbital_keyword)
keyword_prob.append(atom_keyword)
keyword_prob.append(solution_keyword)
keyword_prob.append(basic_keyword)
keyword_prob.append(concentration_keyword)

acid_prob=[]
oxred_prob=[]
precipitaiton_prob=[]
gas_prob=[]
balance_prob=[]
rate_prob=[]
heat_prob=[]
electric_prob=[]
inorganic_prob=[]
organic_prob=[]
periodic_prob=[]
orbital_prob=[]
atom_prob=[]
solution_prob=[]
basic_prob=[]
concentration_prob=[]

chapter_prob = []
for i in prob_lis:
    for j in acid_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            acid_prob.append(prob_lis.index(i)+1)
    for j in oxred_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            oxred_prob.append(prob_lis.index(i)+1)
    for j in precipitaiton_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            precipitaiton_prob.append(prob_lis.index(i)+1)
    for j in  gas_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            gas_prob.append(prob_lis.index(i)+1)
    for j in balance_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            balance_prob.append(prob_lis.index(i)+1)
    for j in rate_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            rate_prob.append(prob_lis.index(i)+1)
    for j in heat_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            heat_prob.append(prob_lis.index(i)+1)
    for j in electric_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            electric_prob.append(prob_lis.index(i)+1)
    for j in inorganic_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            inorganic_prob.append(prob_lis.index(i)+1)
    for j in organic_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            organic_prob.append(prob_lis.index(i)+1)
    for j in periodic_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            periodic_prob.append(prob_lis.index(i)+1)
    for j in orbital_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            orbital_prob.append(prob_lis.index(i)+1)
    for j in atom_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            atom_prob.append(prob_lis.index(i)+1)
    for j in solution_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            solution_prob.append(prob_lis.index(i)+1)
    for j in basic_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            basic_prob.append(prob_lis.index(i)+1)
    for j in concentration_keyword:
        if j in i:
            print(j,f"in prob{prob_lis.index(i)+1}")
            concentration_prob.append(prob_lis.index(i)+1)

chapter_prob = []
chapter_prob.append(acid_prob)
chapter_prob.append(oxred_prob)
chapter_prob.append(precipitaiton_prob)
chapter_prob.append(gas_prob)
chapter_prob.append(balance_prob)
chapter_prob.append(rate_prob)
chapter_prob.append(heat_prob)
chapter_prob.append(electric_prob)
chapter_prob.append(inorganic_prob)
chapter_prob.append(organic_prob)
chapter_prob.append(periodic_prob)
chapter_prob.append(orbital_prob)
chapter_prob.append(atom_prob)
chapter_prob.append(solution_prob)
chapter_prob.append(basic_prob)
chapter_prob.append(concentration_prob)




import docx
cnt=0
for i in chapter_prob:
    tmp = set(i)
    chapter_num = keyword_prob[cnt][0]
    output = sorted(list(tmp))
    print(chapter_num,output)
    cnt+=1


"""
doc_acid = docx.Document()
doc_oxred = docx.Document()
for i in acid_prob:
    doc_acid.add_paragraph(prob_lis[i-1])
for j in oxred_prob:
    doc_oxred.add_paragraph(prob_lis[j-1])

doc_acid.save('./acid/tmp.docx')
doc_oxred.save('./oxred/tmp.docx')
"""